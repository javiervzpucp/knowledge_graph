# -*- coding: utf-8 -*-
"""
Created on Mon Nov 11 16:34:42 2024

@author: jveraz
"""

import os
from langchain_community.graphs import Neo4jGraph
from langchain_openai import ChatOpenAI
from langchain_experimental.graph_transformers import LLMGraphTransformer
from neo4j import GraphDatabase
from langchain_core.documents import Document
import networkx as nx
import json
import matplotlib.pyplot as plt
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document as DocxDocument  # Importar para leer archivos .docx

# Load environment variables from a .env file
load_dotenv()

# Set the OpenAI API key from the .env file
api_key = os.getenv("OPENAI_API_KEY")
url = os.getenv("NEO4J_URI")
neo_user = os.getenv("NEO4J_USERNAME")
neo_pass = os.getenv("NEO4J_PASSWORD")

# Borrar datos del grafo existente en Neo4j
driver = GraphDatabase.driver(url, auth=(neo_user, neo_pass))

# Función para borrar el grafo completo
def clear_graph(tx):
    tx.run("MATCH (n) DETACH DELETE n")

# Ejecutar la función para borrar el grafo
with driver.session() as session:
    session.write_transaction(clear_graph)

# Cierra la conexión
driver.close()

# Configurar la conexión con Neo4j utilizando las credenciales cargadas
graph = Neo4jGraph(
    url=url,
    username=neo_user,
    password=neo_pass
)

# Función para leer todos los archivos .txt, .pdf y .docx de la carpeta "archivos"
def read_files_from_folder(folder_path):
    content = ""
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as file:
                content += file.read() + "\n"
        elif filename.endswith(".pdf"):
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
    return content

# Leer todos los archivos de la carpeta "archivos"
folder_path = "archivos"
lines = read_files_from_folder(folder_path)

# Dividir el texto en fragmentos de 1000 caracteres (ajusta este tamaño si es necesario)
chunk_size = 10000
text_chunks = [lines[i:i + chunk_size] for i in range(0, len(lines), chunk_size)]

# Crear grafo de NetworkX
nx_graph = nx.Graph()

# Procesar cada fragmento individualmente
for i, chunk in enumerate(text_chunks):
    print(f"Procesando fragmento {i + 1} de {len(text_chunks)}...")

    # Crear un documento con el fragmento de texto
    document = Document(page_content=chunk)
  
    llm = ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo")  # Configuración del modelo
    llm_transformer = LLMGraphTransformer(llm=llm)

    graph_documents = llm_transformer.convert_to_graph_documents([document])
    graph.add_graph_documents(
        graph_documents,
        baseEntityLabel=True,
        include_source=False
    )
    
    # Agregar nodos y aristas del fragmento al grafo
    for doc in graph_documents:
        for node in doc.nodes:
            node_id = dict(node)["id"]
            content = dict(node)["type"]
            nx_graph.add_node(node_id, content=content)
        
        for edge in doc.relationships:
            edge_dict = vars(edge) if not isinstance(edge, dict) else edge
            source = str(dict(edge_dict["source"])["id"])
            target = str(dict(edge_dict["target"])["id"])
            tipo = edge_dict["type"]
            nx_graph.add_edge(source, target, content=tipo)

# Configura la conexión a Neo4j
uri = url  # Cambia esto según la URI de tu instancia de Neo4j
user = neo_user  # Cambia esto a tu usuario
password = neo_pass  # Cambia esto a tu contraseña
driver = GraphDatabase.driver(uri, auth=(user, password))

# Función para extraer nodos y relaciones
def get_graph_data(tx):
    nodes = []
    relationships = []
    
    # Obtener nodos
    for record in tx.run("MATCH (n) RETURN n"):
        node = record["n"]
        nodes.append({
            "id": node.id,
            "labels": list(node.labels),
            "properties": dict(node)
        })
    
    # Obtener relaciones
    for record in tx.run("MATCH (a)-[r]->(b) RETURN a, b, r"):
        relationship = record["r"]
        relationships.append({
            "source": record["a"].id,
            "target": record["b"].id,
            "type": relationship.type,
            "properties": dict(relationship)
        })
    
    return {"nodes": nodes, "edges": relationships}

# Extrae los datos del grafo y guárdalos en un archivo JSON
with driver.session() as session:
    graph_data = session.read_transaction(get_graph_data)

with open('paucar_graph.json', 'w') as f:
    json.dump(graph_data, f, indent=4)

# Cierra la conexión a Neo4j
driver.close()
            
# Guardar el grafo en JSON
data = nx.node_link_data(nx_graph)
with open("saved_graph.json", "w") as file:
    json.dump(data, file)
    
# Procesar el componente más grande del grafo
Gcc = sorted(nx.connected_components(nx_graph), key=len, reverse=True)
nx_graph = nx_graph.subgraph(Gcc[0])

# Layout del grafo
pos = nx.kamada_kawai_layout(nx_graph)

# Figura
plt.figure(figsize=(10, 10))

# Dibujar aristas
nx.draw_networkx_edges(nx_graph, pos, edge_color="k", width=0.75)

# Dibujar nodos
nx.draw_networkx_nodes(nx_graph, pos, node_size=500, node_color="gold")

# Dibujar etiquetas
nx.draw_networkx_labels(nx_graph, pos, font_size=7, font_weight="bold")

# Guardar la imagen
plt.savefig("graph.png", format="PNG")
plt.show()
